import logging
import os
import re
import io
from datetime import datetime
from typing import List
import tempfile
import yagmail
from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import FSInputFile, Message, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

def convert_to_pdf(docx_path, pdf_path):
    doc = aw.Document(docx_path)
    doc.save(pdf_path)

# Конфигурация 
API_TOKEN = "8332637912:AAFcZEMZspHLpNU2EQ_tm2crPnwNF3Irk10"
TEMPLATE_PATH = "act.docx"
SMTP_LOGIN = "qxxntxm30@gmail.com"
SMTP_PASSWORD = "urhsobfnzpnscdqn"  # пароль приложения
# пароль приложения
GROUP_ID =-4630725474 # ID группы по умолчанию (можно оставить 0 для отключения)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

bot = Bot(token=API_TOKEN)
dp = Dispatcher(storage=MemoryStorage())

# Состояния 
class ActState(StatesGroup):
    waiting_block = State()
    waiting_description = State()
    waiting_photos = State()
    ask_more = State()
    waiting_destination = State()
    waiting_email = State()
    done = State()

class Damage:
    def __init__(self, block: str, desc: str, photos: List[str]):
        self.block = block
        self.description = desc
        self.photos = photos

def get_main_menu():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📋 Начать акт", callback_data="start_act")]
    ])

def photo_done_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Завершить фото", callback_data="finish_photos")]
    ])

def get_destination_kb():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="📧 На почту", callback_data="send_email"),
            InlineKeyboardButton(text="📱 В группу", callback_data="send_group")
        ],
        [InlineKeyboardButton(text="📧📱 На почту и в группу", callback_data="send_both")],
        [InlineKeyboardButton(text="🚫 Никуда не отправлять", callback_data="send_none")]
    ])

async def ask_destination(message: Message, state: FSMContext):
    await message.answer("Куда отправить акт?", reply_markup=get_destination_kb())
    await state.set_state(ActState.waiting_destination)

# Handlers
@dp.message(Command("start"))
async def cmd_start(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Добро пожаловать!\nНажмите кнопку ниже, чтобы начать составление акта.", reply_markup=get_main_menu())

@dp.callback_query(F.data == "start_act")
async def handle_start_act(callback: CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.answer("📝 Введите № (Блок-Помещение):")
    await state.set_state(ActState.waiting_block)
    await callback.answer()

@dp.message(ActState.waiting_block)
async def get_block(message: Message, state: FSMContext):
    block = message.text.strip()
    if not block:
        await message.answer("⚠️ Введите блок (помещение):")
        return
    await state.update_data(current_block=block)
    await message.answer("Введите описание повреждения:")
    await state.set_state(ActState.waiting_description)

@dp.message(ActState.waiting_description)
async def get_description(message: Message, state: FSMContext):
    desc = message.text.strip()
    if not desc:
        await message.answer("⚠️ Введите описание повреждения:")
        return
    await state.update_data(current_description=desc, current_photos=[])
    await message.answer(
        "📷 Отправьте фото повреждения (можно несколько). После — нажмите кнопку ниже.",
        reply_markup=photo_done_kb()
    )
    await state.set_state(ActState.waiting_photos)

@dp.message(ActState.waiting_photos, F.photo)
async def get_photos(message: Message, state: FSMContext):
    data = await state.get_data()
    photo_list = data.get("current_photos", [])
    if len(photo_list) >= 10:  # MAX_PHOTOS_PER_BLOCK
        await message.answer(f"⚠️ Максимум — 10 фото. Нажмите кнопку «Завершить фото».")
        return
    file_id = message.photo[-1].file_id
    photo_list.append(file_id)
    await state.update_data(current_photos=photo_list)
    await message.answer(f"✅ Фото получено. Всего: {len(photo_list)}.")

@dp.callback_query(F.data == "finish_photos")
async def on_finish_photos(callback: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    photos = data.get("current_photos", [])
    if not photos:
        await callback.answer("⚠️ Сначала отправьте хотя бы одно фото.", show_alert=True)
        return

    damage = Damage(data["current_block"], data["current_description"], photos)
    damages = data.get("damages", [])
    damages.append(damage)
    await state.update_data(
        damages=damages,
        current_block=None,
        current_description=None,
        current_photos=[]
    )

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Добавить помещение", callback_data="add_block")],
        [InlineKeyboardButton(text="✅ Завершить акт", callback_data="finish_act")]
    ])
    await state.set_state(ActState.ask_more)
    await callback.message.answer("Выберите действие:", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data.in_({"add_block", "finish_act"}))
async def handle_next_step(callback: CallbackQuery, state: FSMContext):
    if callback.data == "add_block":
        await callback.message.answer("Введите № (Блок-Помещение):")
        await state.set_state(ActState.waiting_block)
    elif callback.data == "finish_act":
        await ask_destination(callback.message, state)
    await callback.answer()

@dp.callback_query(ActState.waiting_destination, F.data.startswith("send_"))
async def handle_destination(callback: CallbackQuery, state: FSMContext):
    action = callback.data
    await state.update_data(send_action=action)
    
    data = await state.get_data()
    
    if action in ["send_email", "send_both"] and not data.get("email"):
        await callback.message.answer("Введите email для отправки:")
        await state.set_state(ActState.waiting_email)
    else:
        await state.set_state(ActState.done)
        await generate_and_send_act(callback.message, state)
    
    await callback.answer()

@dp.message(ActState.waiting_email)
async def get_email(message: Message, state: FSMContext):
    email = message.text.strip()
    if not re.match(r"^[\w\.-]+@[\w\.-]+\.\w+$", email):
        await message.answer("❌ Неверный email. Повторите ввод:")
        return
    
    await state.update_data(email=email)
    await state.set_state(ActState.done)
    await generate_and_send_act(message, state)

async def generate_and_send_act(message: Message, state: FSMContext):
    data = await state.get_data()
    damages = data.get("damages", [])
    
    if not damages:
        await message.answer("⚠️ Нет данных. Начните заново.")
        await state.clear()
        return

    # Генерация документа
    doc = Document(TEMPLATE_PATH)
    
    # Заполнение шаблона
    for para in doc.paragraphs:
        if "В ходе осмотра были выявлены следующие повреждения" in para.text:
            para.clear()
            run = para.add_run("В ходе осмотра были выявлены следующие повреждения:\n")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            for idx, d in enumerate(damages, 1):
                run = para.add_run(f"{idx}. Блок {d.block} — {d.description}\n")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    # Добавление фотографий
    for d in damages:
        doc.add_page_break()
        doc.add_paragraph(f"Приложение №1 к акту от {datetime.today().strftime('%d.%m.%Y')}")
        doc.add_paragraph(d.block)
        for pid in d.photos:
            try:
                file = await bot.get_file(pid)
                file_data = await bot.download(file.file_id)
                doc.add_picture(io.BytesIO(file_data.read()), width=Inches(4))
            except Exception as e:
                logger.error(f"Ошибка при загрузке фото: {e}")

    # Сохранение в DOCX
    filename = f"акт_{datetime.today().strftime('%d%m%Y_%H%M%S')}.docx"
    temp_path_docx = os.path.join(tempfile.gettempdir(), filename)
    doc.save(temp_path_docx)

    # Отправка DOCX файла пользователю
    await message.answer_document(FSInputFile(temp_path_docx), caption="✅ Акт составлен.")
    
    # Определяем действия для отправки
    action = data.get("send_action", "send_none")
    
    # Отправка в Telegram
    if action in ["send_group", "send_both"] and GROUP_ID != 0:
        try:
            await bot.send_document(
                chat_id=GROUP_ID, 
                document=FSInputFile(temp_path_docx), 
                caption=f"📄 Новый акт дефектовки от @{message.from_user.username}"
            )
            await message.answer(f"✅ Акт отправлен в группу {GROUP_ID}")
        except Exception as e:
            logger.error(f"Ошибка отправки в группу: {e}")
            await message.answer("❌ Не удалось отправить акт в группу")

    # Отправка на почту
    if action in ["send_email", "send_both"] and data.get("email"):
        try:
            yag = yagmail.SMTP(SMTP_LOGIN, SMTP_PASSWORD)
            yag.send(
                to=data["email"],
                subject=f"Акт дефектовки {datetime.today().strftime('%d.%m.%Y')}",
                contents="В приложении акт осмотра.",
                attachments=temp_path_docx  # Изменено на DOCX
            )
            await message.answer(f"📧 Акт отправлен на {data['email']}")
        except Exception as e:
            logger.error(f"Ошибка отправки email: {e}")
            await message.answer(f"❌ Не удалось отправить на {data['email']}")

    await state.clear()
    await message.answer("🏠 Возврат в главное меню:", reply_markup=get_main_menu())


if __name__ == "__main__":
    import asyncio
    asyncio.run(dp.start_polling(bot))
